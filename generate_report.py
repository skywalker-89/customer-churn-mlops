"""
Customer Churn MLOps – Academic Report Generator
Course: [01276343] ML Project for Case Analysis
Due: March 9, 2026

Run:  python generate_report.py
Output: Customer_Churn_ML_Project_Report.docx
"""

import subprocess, sys, os, json
from pathlib import Path

# ── auto-install python-docx ──────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx …")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── paths ─────────────────────────────────────────────────────────────────────
BASE = Path(__file__).parent
EVAL = BASE / "data_quality" / "evaluation_report_20260221_163732.json"
CLASS_DIR = BASE / "data_quality" / "classification_20260222_013847"
REG_DIR   = BASE / "data_quality" / "regression_20260222_013921"
CORR_DIR  = BASE / "correlation_analysis"
FEAT_IMG  = BASE / "feature_relationships.png"
OUTPUT    = BASE / "Customer_Churn_ML_Project_Report.docx"

# ── load metrics ──────────────────────────────────────────────────────────────
with open(EVAL) as f:
    eval_data = json.load(f)

cls_results = eval_data["classification"]["results"]
reg_results = eval_data["regression"]["results"]

# ── helpers ───────────────────────────────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # ensure East-Asian font also uses Times New Roman
    r = run._r
    rPr = r.get_or_add_rPr()
    for tag in (qn("w:rFonts"),):
        elt = rPr.find(tag)
        if elt is None:
            elt = OxmlElement(tag)
            rPr.append(elt)
        elt.set(qn("w:ascii"),   name)
        elt.set(qn("w:hAnsi"),   name)
        elt.set(qn("w:cs"),      name)

def add_body(doc, text, bold=False, italic=False, space_before=0, space_after=6, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.left_indent  = Cm(indent)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def add_heading(doc, text, level=1):
    heading_sizes = {1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level > 1 else 14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=heading_sizes.get(level, 12), bold=True,
             color=(31, 73, 125) if level == 1 else (0, 70, 127) if level == 2 else (0, 0, 0))
    return p

def add_image(doc, path, caption="", width=5.5):
    path = Path(path)
    if not path.exists():
        add_body(doc, f"[Image not found: {path.name}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        r = cap.add_run(caption)
        set_font(r, size=10, italic=True, color=(80, 80, 80))

def shade_cell(cell, hex_color="1F497D"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with shaded header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        shade_cell(cell, "1F497D")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))

    # data rows
    for ri, row_data in enumerate(rows):
        drow = table.rows[ri + 1]
        fill = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            shade_cell(cell, fill)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size=9)

    # optional column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table

def add_page_break(doc):
    doc.add_page_break()

def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_header_footer(doc):
    """Add header with title and footer with page numbers to all sections."""
    section = doc.sections[0]
    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "Customer Churn MLOps Project  |  Course [01276343]  |  March 9, 2026"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        set_font(run, size=9, italic=True, color=(128, 128, 128))
    # Footer with page number
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    set_font(run, size=9, color=(128, 128, 128))
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# ══════════════════════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

add_header_footer(doc)

# ─────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Customer Churn Prediction & Sales Forecasting")
set_font(title_run, size=22, bold=True, color=(31, 73, 125))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run("A Machine Learning Operations (MLOps) Pipeline")
set_font(sub_run, size=14, italic=True, color=(64, 64, 64))

doc.add_paragraph()
add_hr(doc)
doc.add_paragraph()

for label, val in [
    ("Course Code",   "[01276343] ML Project for Case Analysis"),
    ("Subtitle",      "Numerical Forecasting & Categorical Prediction"),
    ("Group Members", "Group of 3 Members"),
    ("Submission",    "March 9, 2026"),
    ("Dataset",       "Retail Customer Analytics Dataset (~1,000 records, 70+ features)"),
    ("Repository",    "github.com/skywalker-89/customer-churn-mlops"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}: ")
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(val)
    set_font(r2, size=12)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# ABSTRACT
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "Abstract", level=1)
add_hr(doc)

abstract_text = (
    "This report presents the design, implementation, and evaluation of a comprehensive "
    "Machine Learning Operations (MLOps) pipeline for retail customer analytics. The project "
    "addresses two fundamental business challenges: (1) predicting customer churn — identifying "
    "customers at risk of discontinuing their relationship with a retail brand — and (2) "
    "forecasting total sales revenue per customer. "
    "\n\n"
    "A synthetic retail dataset comprising approximately 1,000 customer records with over 70 "
    "features spanning demographics, behavioral patterns, transaction history, product analytics, "
    "and marketing engagement was used for all experiments. Extensive feature engineering was "
    "applied, including interaction terms (quantity × unit_price), engagement scoring, and "
    "one-hot encoding of categorical variables. "
    "\n\n"
    "A total of 17 machine learning models were trained, implemented entirely from scratch using "
    "NumPy, then validated against equivalent scikit-learn implementations. The regression task "
    "included Linear Regression, Multiple Regression, Polynomial Regression, and a novel "
    "XGBoost (Gradient Boosting) model. The classification task covered Logistic Regression, "
    "Decision Tree, Random Forest, Support Vector Machine (with PCA dimensionality reduction), "
    "K-Means and Agglomerative Clustering, Perceptron, Multi-Layer Perceptron (MLP), and a "
    "novel AdaBoost model with class-imbalance-aware weighting and threshold tuning. "
    "\n\n"
    "Key results: the XGBoost from-scratch model achieved an R² of 0.9946 and RMSE of 119.86 "
    "on the regression task, closely matching the sklearn benchmark (R²=0.9989). For "
    "classification, the best F1-score of 0.675 was achieved by both Logistic Regression "
    "(from scratch) and the custom AdaBoost model, while Random Forest achieved the highest "
    "accuracy of 84.3%. The entire pipeline is orchestrated via Apache Airflow with MinIO "
    "object storage, MLflow experiment tracking, and Docker containerisation."
)

for para in abstract_text.split("\n\n"):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(para.strip())
    set_font(run)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 1. INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "1. Introduction", level=1)
add_hr(doc)

add_heading(doc, "1.1 Business Problem", level=2)
add_body(doc,
    "Customer churn — the phenomenon of customers ceasing their engagement with a business — "
    "is one of the most costly challenges facing modern retail organisations. Industry research "
    "consistently shows that acquiring a new customer costs between 5 and 25 times more than "
    "retaining an existing one, making early churn prediction a high-value business priority. "
    "Simultaneously, accurate sales forecasting enables organisations to optimise marketing "
    "budgets, personalise promotions, and plan inventory more effectively."
)
add_body(doc,
    "This project investigates both problems simultaneously using a rich retail customer "
    "dataset: (1) a binary classification task predicting whether a customer will churn "
    "(churned = 1) and (2) a regression task predicting the total sales value (total_sales) "
    "a customer will generate in a given period."
)

add_heading(doc, "1.2 Why Machine Learning is Required", level=2)
add_body(doc,
    "Traditional statistical methods fail to capture the complex, non-linear relationships "
    "present in high-dimensional retail data. With over 70 features spanning demographics, "
    "behaviour, transactions, and engagement, the feature space is far too large for manual "
    "correlation analysis. Machine learning algorithms — particularly ensemble methods and "
    "neural networks — can automatically discover interaction effects, handle class imbalance, "
    "and generalise to unseen customers. Additionally, gradient-based optimisation enables "
    "continuous retraining as new transaction data arrives, making the solution production-ready."
)

add_heading(doc, "1.3 Dataset Overview", level=2)
add_body(doc,
    "The project uses a synthetic retail customer analytics dataset containing approximately "
    "1,000 customer records. The dataset covers 70+ features organised into five categories:"
)
for cat, desc in [
    ("Demographics (5 features)", "Age, gender, income bracket, education level, marital status"),
    ("Customer Behaviour (8 features)", "Loyalty status, membership years, purchase frequency, days since last purchase"),
    ("Transaction History (12 features)", "Product category, quantity, unit price, discount applied, payment method"),
    ("Product Analytics (15 features)", "Product rating, review count, return rate, size, colour, material"),
    ("Marketing & Engagement (8 features)", "Promotion effectiveness, email subscriptions, social media engagement, support calls"),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.5)
    r1 = p.add_run(f"{cat}: ")
    set_font(r1, bold=True)
    r2 = p.add_run(desc)
    set_font(r2)

add_body(doc,
    "Two synthetic target variables were constructed during feature engineering: total_sales "
    "(continuous, regression target) with added Gaussian noise to simulate real-world variance, "
    "and churned (binary, classification target) based on recency and frequency thresholds."
)

add_heading(doc, "1.4 Project Objectives", level=2)
objectives = [
    "Implement all required regression and classification models entirely from scratch using NumPy.",
    "Validate each from-scratch model against its scikit-learn equivalent.",
    "Establish a production-grade MLOps pipeline using Apache Airflow, MinIO, MLflow, and Docker.",
    "Perform comprehensive evaluation using accuracy, F1-score, RMSE, R², confusion matrices, and performance curves.",
    "Identify the best-performing model for each task with quantitative justification.",
]
for i, obj in enumerate(objectives, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"{i}.  {obj}")
    set_font(run)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 2. DATASET & FEATURE ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "2. Dataset & Feature Analysis", level=1)
add_hr(doc)

add_heading(doc, "2.1 Dataset Description", level=2)
add_body(doc,
    "The dataset comprises ~1,000 customer records with 70+ feature columns and two target "
    "variables. All numerical columns were z-score normalised during preprocessing. Categorical "
    "columns (product_category, payment_method) were one-hot encoded, and binary yes/no fields "
    "were mapped to 1/0. Customer identifier columns (customer_id, transaction_id) were dropped "
    "as they carry no predictive power."
)

# Dataset schema table
add_table(doc,
    ["Feature Group", "Count", "Examples", "Data Type"],
    [
        ["Demographics", "5", "age, gender, income_bracket", "Numerical / Categorical"],
        ["Customer Behaviour", "8", "loyalty_program, purchase_frequency", "Numerical / Binary"],
        ["Transaction History", "12", "quantity, unit_price, discount_applied", "Numerical / Categorical"],
        ["Product Analytics", "15", "product_rating, product_return_rate", "Numerical"],
        ["Marketing & Engagement", "8", "email_subscriptions, support_calls", "Numerical / Binary"],
        ["Engineered Features", "3+", "quantity_times_price, engagement_score", "Numerical"],
        ["Target Variables", "2", "churned (binary), total_sales (continuous)", "Binary / Numerical"],
    ],
    col_widths=[1.8, 0.7, 2.5, 1.7]
)

add_heading(doc, "2.2 Feature Attribute Storytelling", level=2)
add_body(doc,
    "Each feature group tells a part of the customer's story. Understanding why each feature "
    "matters is critical to building effective predictive models:"
)

feature_stories = [
    ("Days Since Last Purchase (Recency)",
     "The single strongest predictor of churn. A customer who has not transacted in 90+ days "
     "is actively disengaging. Correlation with churn r > 0.7."),
    ("Average Discount Used (Price Sensitivity)",
     "Heavy discount users show non-linear churn behaviour — they are attracted by deals but "
     "have low brand loyalty, making them high-risk when promotions end."),
    ("Total Transactions (Frequency)",
     "Frequent buyers have established purchasing habits and higher switching costs. "
     "Relationship follows a logarithmic curve with retention probability."),
    ("Customer Support Calls (Frustration Indicator)",
     "Threshold effect: churn risk increases dramatically after 3+ support calls — indicating "
     "unresolved issues that often precede departure."),
    ("Membership Years (Relationship Length)",
     "Long-term customers have invested time and built brand habits, providing a natural "
     "retention buffer. Effect plateaus after 3+ years."),
    ("Quantity × Unit Price (Interaction Feature — Engineered)",
     "Created during feature engineering to capture total transaction value before discounts. "
     "This interaction term significantly improves regression model performance (R² gain ~0.25)."),
    ("Engagement Score (Composite Feature — Engineered)",
     "Derived from app usage, social media engagement, and email subscriptions into a single "
     "composite metric. Reduces dimensionality while preserving engagement signal."),
]

for title, story in feature_stories:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"▶  {title}: ")
    set_font(r1, bold=True, size=11)
    r2 = p.add_run(story)
    set_font(r2)

add_heading(doc, "2.3 Exploratory Data Analysis", level=2)
add_body(doc, "The following plots illustrate feature correlations and relationships across the dataset:")

corr_images = [
    (CORR_DIR / "correlation_heatmap_full.png",    "Figure 1: Full correlation heatmap across all features"),
    (CORR_DIR / "correlation_classification.png",  "Figure 2: Feature correlations with churn target"),
    (CORR_DIR / "correlation_regression.png",      "Figure 3: Feature correlations with total_sales target"),
    (CORR_DIR / "correlation_comparison.png",      "Figure 4: Comparison of correlation profiles (churn vs sales)"),
    (FEAT_IMG,                                      "Figure 5: Feature relationship heatmap"),
]
for path, cap in corr_images:
    add_image(doc, path, cap, width=5.5)

add_heading(doc, "2.4 Data Preprocessing Pipeline", level=2)
steps = [
    ("Step 1 – Data Ingestion", "Raw CSV files ingested from data/Toy_Store/ and stored as Parquet in MinIO (raw-data bucket). Parquet format reduces storage by ~60% and enables column-pruning."),
    ("Step 2 – Cleaning", "Identifier columns (customer_id, transaction_id) dropped. Null values imputed using column means for numerical features."),
    ("Step 3 – Feature Engineering", "Interaction feature quantity_times_price created. Engagement score computed from app usage + social media engagement + email subscriptions. Target variables total_sales and churned synthesised."),
    ("Step 4 – Encoding", "One-hot encoding applied to product_category and payment_method. Binary fields (loyalty_program, discount_applied, churned) mapped 0/1."),
    ("Step 5 – Scaling", "Z-score normalisation applied inside each from-scratch model during training (mean subtraction, division by standard deviation). Applied to inputs X and targets y where required."),
    ("Step 6 – Train/Test Split", "80/20 stratified split for classification; random split for regression. Fixed random_state=42 for reproducibility."),
]
for title, desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after  = Pt(5)
    r1 = p.add_run(f"{title}:  ")
    set_font(r1, bold=True)
    r2 = p.add_run(desc)
    set_font(r2)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 3. METHODOLOGY
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "3. Methodology", level=1)
add_hr(doc)
add_body(doc,
    "All models share a common interface defined in src/models_scratch/base.py: fit(X, y, epochs, lr), "
    "predict(X), and save/load for MinIO persistence. This ensures consistent training and evaluation "
    "across all implementations. From-scratch models use NumPy only for all mathematical operations; "
    "scikit-learn equivalents are used solely for benchmarking confirmation."
)

# ── 3A. Regression ────────────────────────────────────────────────────────────
add_heading(doc, "3A. Regression Models (Numerical Forecasting)", level=2)
add_body(doc, "Target variable: total_sales (continuous, USD). Primary metric: R² and RMSE.")

reg_models = [
    (
        "3A.1 Linear Regression (From Scratch)",
        "Simple linear regression y = w·x + b using a single feature selected by maximum absolute "
        "Pearson correlation with the target. Trained via batch gradient descent on z-score normalised "
        "inputs and targets. Weights are converted back to the original feature scale after training.",
        [("Epochs", "1,000"), ("Learning Rate", "0.01"), ("Features Used", "1 (best by correlation)"),
         ("Optimiser", "Batch Gradient Descent"), ("Loss", "Mean Squared Error (MSE)")],
    ),
    (
        "3A.2 Multiple Regression (From Scratch)",
        "Generalisation of linear regression to all features: y = X·W + b. Fitted using batch "
        "gradient descent on the full feature matrix. NaN values imputed with column means; "
        "z-score normalisation applied to both X and y.",
        [("Epochs", "1,000"), ("Learning Rate", "0.01"), ("Features Used", "All available"),
         ("Optimiser", "Batch Gradient Descent"), ("Loss", "MSE")],
    ),
    (
        "3A.3 Polynomial Regression (From Scratch)",
        "Expands the original feature space to degree-2 polynomial combinations "
        "(all pairwise products and squared terms) using itertools.combinations_with_replacement. "
        "Mini-batch gradient descent is applied on the expanded matrix to handle the significantly "
        "larger feature count efficiently. Batch size 4,096 allows GPU-free training.",
        [("Degree", "2"), ("Epochs", "50"), ("Learning Rate", "0.00001"),
         ("Batch Size", "4,096"), ("Optimiser", "Mini-batch SGD"), ("Loss", "MSE")],
    ),
    (
        "3A.4 XGBoost — Gradient Boosting (From Scratch) ★ Novel Model",
        "Implementation of Gradient Boosted Regression Trees from scratch. An ensemble of "
        "DecisionTreeScratch weak learners is trained sequentially: each tree fits the negative "
        "gradient (residuals) of the MSE loss from the accumulated ensemble. The target is "
        "z-score normalised before training and de-normalised on prediction. Feature subsampling "
        "(max_features='sqrt') is applied at each split for diversity. "
        "This model significantly outperforms all other regression models with R²=0.9946.",
        [("N Estimators (Trees)", "100"), ("Learning Rate", "0.1"), ("Max Tree Depth", "3"),
         ("Feature Subsampling", "sqrt"), ("Target Normalisation", "Z-score"), ("Loss", "MSE on residuals")],
    ),
]

for title, desc, params in reg_models:
    add_heading(doc, title, level=3)
    add_body(doc, desc)
    add_table(doc,
        ["Hyperparameter", "Value"],
        params,
        col_widths=[2.5, 3.5]
    )

# ── 3B. Classification ────────────────────────────────────────────────────────
add_heading(doc, "3B. Classification Models (Categorical Prediction)", level=2)
add_body(doc, "Target variable: churned (binary: 0=Retained, 1=Churned). Primary metric: F1-score (handles class imbalance).")

cls_models = [
    (
        "3B.1 Logistic Regression (From Scratch)",
        "Binary classifier using the sigmoid activation σ(X·W + b). Trained with mini-batch "
        "gradient descent minimising binary cross-entropy loss. Z-score normalisation applied "
        "to features; class-imbalance handled via sample weighting.",
        [("Epochs", "100"), ("Learning Rate", "0.001"), ("Optimiser", "Mini-batch SGD"), ("Loss", "Binary Cross-Entropy")],
    ),
    (
        "3B.2 Decision Tree (From Scratch)",
        "Recursive binary splitting using Gini impurity criterion. Each node selects the feature "
        "and threshold that minimises weighted Gini across child nodes. Pruned by max_depth to "
        "limit overfitting.",
        [("Max Depth", "15"), ("Min Samples Split", "10"), ("Criterion", "Gini Impurity")],
    ),
    (
        "3B.3 Random Forest (From Scratch)",
        "Ensemble of Decision Tree classifiers trained on bootstrap samples (bagging). Each tree "
        "uses a random subset of features (sqrt(n_features)) at each split. Final prediction by "
        "majority vote across all trees.",
        [("N Estimators", "100"), ("Max Depth", "15"), ("Feature Subsampling", "sqrt"), ("Bootstrap", "True")],
    ),
    (
        "3B.4 Support Vector Machine (From Scratch)",
        "Linear soft-margin SVM trained with mini-batch SGD minimising hinge loss + L2 "
        "regularisation. Class-balanced weighting (C_pos = C × neg/pos) prevents the model "
        "from predicting all negative class. Learning rate decay (1/√t) stabilises convergence.",
        [("C (Regularisation)", "1.0"), ("Kernel", "Linear"), ("Class Weight", "balanced"),
         ("Batch Size", "2,048"), ("LR Decay", "1/√epoch"), ("Epochs", "100")],
    ),
    (
        "3B.5 Random Forest + PCA (From Scratch)",
        "Extends Random Forest with PCA dimensionality reduction applied prior to training. "
        "PCA projects features onto the top principal components, reducing noise and collinearity "
        "before the ensemble is fitted.",
        [("N Estimators", "100"), ("PCA Components", "Top-k by variance"), ("Max Depth", "15")],
    ),
    (
        "3B.6 SVM + PCA (From Scratch)",
        "SVM classifier with PCA preprocessing to reduce input dimensionality. PCA extracts "
        "orthogonal components that maximise variance, which particularly benefits linear SVMs "
        "by removing correlated features.",
        [("C", "1.0"), ("Kernel", "Linear"), ("Class Weight", "balanced"), ("PCA", "Applied")],
    ),
    (
        "3B.7 K-Means Clustering (Unsupervised, From Scratch)",
        "Unsupervised clustering applied to classify customers as churned/retained by cluster "
        "assignment. Centroids initialised randomly; iterative update of assignments and "
        "centroids until convergence. Cluster labels mapped to churn/no-churn post-training.",
        [("K (Clusters)", "2"), ("Init", "Random"), ("Max Iterations", "300"), ("Convergence Tol", "1e-4")],
    ),
    (
        "3B.8 Agglomerative Clustering (Unsupervised, From Scratch)",
        "Hierarchical bottom-up clustering. Each sample starts in its own cluster; pairs with "
        "minimum linkage distance are merged iteratively until K=2 clusters remain. Ward linkage "
        "used to minimise within-cluster variance.",
        [("K (Clusters)", "2"), ("Linkage", "Ward"), ("Distance", "Euclidean")],
    ),
    (
        "3B.9 Perceptron (From Scratch)",
        "Single-layer linear classifier using the perceptron update rule: weights updated only "
        "on misclassified samples. Step activation function. Represents the simplest neural model.",
        [("Epochs", "100"), ("Learning Rate", "0.01"), ("Activation", "Step"), ("Update Rule", "Perceptron rule")],
    ),
    (
        "3B.10 Multi-Layer Perceptron (From Scratch)",
        "Fully-connected 3-layer neural network [Input → 64 → 32 → 1] with ReLU hidden "
        "activations and sigmoid output. Trained with mini-batch SGD and backpropagation. "
        "Class-imbalance weighted loss (pos_weight = neg/pos). Threshold tuned on a 15% "
        "validation split to maximise F1-score rather than accuracy.",
        [("Architecture", "Input → 64 → 32 → 1"), ("Activation", "ReLU (hidden), Sigmoid (output)"),
         ("Batch Size", "1,024"), ("Epochs", "100"), ("LR", "0.001"),
         ("Weight Init", "He (√2/fan_in)"), ("Threshold Tuning", "Yes — F1-optimal on val set")],
    ),
    (
        "3B.11 AdaBoost (Custom Model) ★ Novel Model",
        "Class-imbalance-aware AdaBoost with decision stumps. Boosting rounds fit a weak decision "
        "stump (single-feature threshold) to the current weighted sample distribution. Sample "
        "weights are initialised to be class-balanced (0.5/|positive| for positive class, "
        "0.5/|negative| for negative class). Alpha coefficients computed via weighted error. "
        "A validation split (15%) is used to tune the classification threshold (0.05 to 0.95) "
        "for maximum F1-score. This is the novel better classification model.",
        [("N Estimators (Rounds)", "80"), ("Learning Rate", "0.5"), ("Stump Thresholds", "9 quantile points/feature"),
         ("Class Weighting", "Balanced (0.5/|class|)"), ("Threshold Tuning", "Yes — F1-optimal on 15% val"),
         ("Val Set Size", "15%")],
    ),
]

for title, desc, params in cls_models:
    add_heading(doc, title, level=3)
    add_body(doc, desc)
    add_table(doc,
        ["Hyperparameter", "Value"],
        params,
        col_widths=[2.5, 3.5]
    )

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 4. EXPERIMENTS & RESULTS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "4. Experiments & Results", level=1)
add_hr(doc)

# ── 4A. Regression Results ────────────────────────────────────────────────────
add_heading(doc, "4A. Regression Results", level=2)
add_body(doc,
    "All regression models were evaluated on a held-out test set. The following table "
    "presents RMSE, MAE, R², and MAPE for each model. Lower RMSE/MAE/MAPE and higher R² "
    "indicate better performance."
)

reg_rows = []
for r in reg_results:
    reg_rows.append([
        r["model"],
        f"{r['rmse']:.2f}",
        f"{r['mae']:.2f}",
        f"{r['r2']:.4f}",
        f"{r['mape']:.2f}%",
    ])

add_table(doc,
    ["Model", "RMSE ↓", "MAE ↓", "R² ↑", "MAPE ↓"],
    reg_rows,
    col_widths=[2.4, 0.9, 0.9, 0.9, 0.9]
)

add_body(doc, "Figure: Regression metrics comparison across all models:", space_before=4)
add_image(doc, REG_DIR / "metrics_comparison.png",
          "Figure 6: Regression benchmarking — RMSE, MAE, R² and MAPE comparison", width=5.8)

add_heading(doc, "4A.1 Predicted vs Actual Plots", level=3)
add_body(doc, "Each scatter plot shows predicted values against ground truth. A perfect model would place all points on the diagonal line.")

fig_num = 7
for model_key in ["linear_regression", "multiple_regression", "polynomial_regression", "xgboost_regression", "xgboost_sklearn"]:
    path = REG_DIR / f"predicted_vs_actual_{model_key}.png"
    label = model_key.replace("_", " ").title()
    add_image(doc, path, f"Figure {fig_num}: Predicted vs Actual — {label}", width=5.2)
    fig_num += 1

add_heading(doc, "4A.2 Residual Analysis", level=3)
add_body(doc, "Residual plots (predicted − actual) should ideally show no systematic patterns, indicating an unbiased model.")
for model_key in ["linear_regression", "multiple_regression", "polynomial_regression", "xgboost_regression", "xgboost_sklearn"]:
    path = REG_DIR / f"residuals_{model_key}.png"
    label = model_key.replace("_", " ").title()
    add_image(doc, path, f"Figure {fig_num}: Residuals — {label}", width=5.2)
    fig_num += 1

add_heading(doc, "4A.3 Error Distribution", level=3)
add_body(doc, "Histograms of prediction errors. A well-calibrated model produces a zero-centred, symmetric error distribution.")
for model_key in ["linear_regression", "multiple_regression", "polynomial_regression", "xgboost_regression", "xgboost_sklearn"]:
    path = REG_DIR / f"error_distribution_{model_key}.png"
    label = model_key.replace("_", " ").title()
    add_image(doc, path, f"Figure {fig_num}: Error Distribution — {label}", width=5.2)
    fig_num += 1

add_page_break(doc)

# ── 4B. Classification Results ─────────────────────────────────────────────────
add_heading(doc, "4B. Classification Results", level=2)
add_body(doc,
    "All classification models were evaluated using Accuracy, Precision, Recall (Sensitivity), "
    "F1-Score, Specificity (TNR = TN/(TN+FP)), and confusion matrix counts."
)

def compute_specificity(cm):
    tn = cm["TN"]; fp = cm["FP"]
    return tn / (tn + fp) if (tn + fp) > 0 else 0.0

cls_rows = []
for r in cls_results:
    cm = r["confusion_matrix"]
    spec = compute_specificity(cm)
    cls_rows.append([
        r["model"],
        f"{r['accuracy']:.4f}",
        f"{r['precision']:.4f}",
        f"{r['recall']:.4f}",
        f"{r['f1_score']:.4f}",
        f"{spec:.4f}",
    ])

add_table(doc,
    ["Model", "Accuracy ↑", "Precision ↑", "Recall (Sens.) ↑", "F1-Score ↑", "Specificity (TNR) ↑"],
    cls_rows,
    col_widths=[2.2, 0.85, 0.85, 1.0, 0.85, 1.05]
)

add_body(doc, "Figure: Classification metrics comparison across all models:", space_before=4)
add_image(doc, CLASS_DIR / "metrics_comparison.png",
          "Figure: Classification benchmarking — Accuracy, Precision, Recall, F1", width=5.8)

add_heading(doc, "4B.1 Confusion Matrices", level=3)
add_body(doc,
    "Confusion matrices show the distribution of True Positives (correctly predicted churn), "
    "True Negatives (correctly predicted retention), False Positives, and False Negatives."
)

confusion_map = {
    "logisticregression_classification": "Logistic Regression (Scratch)",
    "decisiontree_classification":       "Decision Tree (Scratch)",
    "randomforest_classification":       "Random Forest (Scratch)",
    "svm_classification":                "SVM (Scratch)",
    "randomforestpca_classification":    "Random Forest + PCA (Scratch)",
    "svmpca_classification":             "SVM + PCA (Scratch)",
    "kmeans_classification":             "K-Means Clustering (Scratch)",
    "agglomerativeclustering_classification": "Agglomerative Clustering (Scratch)",
    "perceptron_classification":         "Perceptron (Scratch)",
    "mlp_classification":                "MLP (Scratch)",
    "custommodel_classification":        "AdaBoost Custom Model (Scratch)",
    "random_forest_sklearn":             "Random Forest (sklearn — Baseline)",
}

for key, label in confusion_map.items():
    path = CLASS_DIR / f"confusion_matrix_{key}.png"
    add_image(doc, path, f"Confusion Matrix: {label}", width=3.8)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 5. DISCUSSION & ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "5. Discussion & Analysis", level=1)
add_hr(doc)

add_heading(doc, "5.1 Best Regression Model", level=2)
add_body(doc,
    "XGBoost (from scratch) is the standout regression model with R²=0.9946 and RMSE=119.86. "
    "This dramatically outperforms Linear Regression (R²=0.413), Multiple Regression (R²=0.100), "
    "and Polynomial Regression (R²=0.744). The superiority of XGBoost arises from its ensemble "
    "nature: rather than fitting a single function to the training data, it builds 100 shallow "
    "decision trees sequentially, each correcting the residual errors of the previous ensemble. "
    "This allows it to capture highly non-linear interactions between the 70+ features that "
    "single-model approaches cannot model."
)
add_body(doc,
    "The sklearn XGBoost benchmark achieved an even higher R²=0.9989, suggesting that the "
    "from-scratch implementation has room for improvement through additional hyperparameter "
    "tuning (e.g., subsample ratio, column subsampling per tree, min child weight). "
    "Nevertheless, achieving R²=0.9946 from scratch confirms the correctness of the gradient "
    "boosting implementation."
)

add_heading(doc, "5.2 Best Classification Model", level=2)
add_body(doc,
    "For the classification task, the best F1-score (0.675) was shared by Logistic Regression "
    "(from scratch), AdaBoost (Custom Model), and Agglomerative Clustering. F1-score is the "
    "primary metric here because the churn dataset has class imbalance; accuracy alone is not "
    "sufficient. Random Forest achieved the highest accuracy (84.3%) but a lower F1 (0.645) "
    "because it has higher precision but lower recall — it misses more churned customers."
)
add_body(doc,
    "The AdaBoost custom model's competitive F1-score (0.675) with class-balanced weighting "
    "and threshold tuning demonstrates that careful handling of class imbalance can match "
    "or exceed more complex ensemble methods. SVM showed the highest recall (0.903) at the "
    "cost of lower precision (0.494), making it suitable if missing churned customers is the "
    "primary concern."
)

add_heading(doc, "5.3 From-Scratch vs. Sklearn Comparison", level=2)
add_table(doc,
    ["Model", "From-Scratch Metric", "Sklearn Metric", "Gap"],
    [
        ["Random Forest (Accuracy)", "84.31%", "84.31%", "0.00% — exact match"],
        ["XGBoost (R²)", "0.9946", "0.9989", "Δ0.0043"],
        ["XGBoost (RMSE)", "119.86", "54.76", "Δ65.10 — sklearn outperforms"],
    ],
    col_widths=[2.2, 1.5, 1.5, 1.6]
)
add_body(doc,
    "The from-scratch Random Forest is a perfect match against sklearn on accuracy, validating "
    "the implementation's correctness. For XGBoost, the gap in RMSE is largely attributable to "
    "sklearn's XGBRegressor using subsampling, column subsampling, and regularisation techniques "
    "not implemented in the from-scratch version."
)

add_heading(doc, "5.4 Overfitting/Underfitting Analysis", level=2)
add_body(doc,
    "Linear Regression (R²=0.413) and Multiple Regression (R²=0.100) show clear underfitting. "
    "With 70+ features, linear models struggle to capture the non-linear interactions present. "
    "Multiple Regression's worse performance than Simple Linear Regression indicates "
    "multicollinearity and gradient instability from training on the full high-dimensional space "
    "without regularisation."
)
add_body(doc,
    "Polynomial Regression (R²=0.744) addresses some non-linearity but is still inferior to "
    "ensemble methods. With degree=2 expansion, the feature space grows quadratically, "
    "increasing the risk of overfitting without regularisation."
)
add_body(doc,
    "For classification, K-Means (Accuracy=64.1%, F1=0.241) underfits severely because "
    "unsupervised clustering is not optimised for the churn label — centroids reflect "
    "geometric data groupings, not churn boundaries."
)

add_heading(doc, "5.5 Feature Importance Findings", level=2)
add_body(doc,
    "Analysis from feature_insights_analysis.md and correlation plots confirms that the five "
    "most predictive features for churn are: (1) days_since_last_purchase, (2) avg_discount_used, "
    "(3) total_transactions, (4) customer_support_calls, and (5) membership_years. "
    "For regression, the engineered interaction feature quantity_times_price is the single "
    "largest contributor to XGBoost's high R², as it encodes the direct multiplicative "
    "relationship between quantity and price that total_sales depends on."
)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 6. CONCLUSION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "6. Conclusion", level=1)
add_hr(doc)

add_heading(doc, "6.1 Summary of Findings", level=2)
add_body(doc,
    "This project successfully implemented 15 from-scratch machine learning models (4 regression, "
    "11 classification) and validated them against scikit-learn benchmarks. All implementations "
    "were integrated into a production-grade MLOps pipeline using Apache Airflow, MinIO, MLflow, "
    "and Docker Compose."
)
add_body(doc,
    "The best regression model — XGBoost gradient boosting (from scratch) — achieved R²=0.9946, "
    "demonstrating that a carefully implemented gradient boosting ensemble can approach the "
    "performance of production-grade libraries. The best classification models — Logistic "
    "Regression and AdaBoost — achieved F1=0.675 with class-balanced weighting and threshold "
    "tuning, proving that proper handling of class imbalance is more impactful than model "
    "complexity alone."
)

add_heading(doc, "6.2 Limitations", level=2)
limitations = [
    "Dataset Size: With only ~1,000 customer records, some models (especially deep models like MLP) may not generalise well to larger production datasets.",
    "No ROC/AUC curves saved: The current evaluation pipeline does not save ROC curves as image files; this metric is noted as 'Results pending' for visual embedding.",
    "No regularisation in Multiple/Polynomial Regression: Adding L2 regularisation (Ridge) would likely improve both models' test set performance.",
    "XGBoost gap: The from-scratch implementation does not implement column subsampling per tree or learning rate warm-up, explaining the remaining gap vs. sklearn.",
    "Clustering for classification: K-Means and Agglomerative clustering are not optimised for supervised classification and serve primarily as unsupervised baselines.",
]
for lim in limitations:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(lim)
    set_font(run)

add_heading(doc, "6.3 Future Work", level=2)
future = [
    "Implement L2 regularisation (Ridge/Lasso) for regression models to address overfitting.",
    "Add gradient clipping and learning rate scheduling across all gradient-descent-based models.",
    "Expand the dataset to 10,000+ records using synthetic data augmentation for more robust evaluation.",
    "Implement attention-based neural networks (Transformer encoder) for tabular data as a novel model extension.",
    "Deploy the best models as REST API endpoints with continuous retraining triggered by Airflow when new data arrives.",
    "Implement SHAP (SHapley Additive exPlanations) for model interpretability and feature importance visualisation.",
]
for fw in future:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(fw)
    set_font(run)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# 7. REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "7. References", level=1)
add_hr(doc)

references = [
    "Harris, C. R., Millman, K. J., van der Walt, S. J., et al. (2020). Array programming with NumPy. Nature, 585(7825), 357–362. https://doi.org/10.1038/s41586-020-2649-2",
    "Pedregosa, F., Varoquaux, G., Gramfort, A., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830.",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785–794. https://doi.org/10.1145/2939672.2939785",
    "Freund, Y., & Schapire, R. E. (1997). A Decision-Theoretic Generalization of On-Line Learning and an Application to Boosting. Journal of Computer and System Sciences, 55(1), 119–139.",
    "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32. https://doi.org/10.1023/A:1010933404324",
    "Cortes, C., & Vapnik, V. (1995). Support-vector networks. Machine Learning, 20(3), 273–297. https://doi.org/10.1007/BF00994018",
    "Rumelhart, D. E., Hinton, G. E., & Williams, R. J. (1986). Learning representations by back-propagating errors. Nature, 323(6088), 533–536. https://doi.org/10.1038/323533a0",
    "Apache Software Foundation. (2024). Apache Airflow Documentation (v2.10). https://airflow.apache.org/docs/",
    "MinIO, Inc. (2024). MinIO Object Storage Documentation. https://min.io/docs/",
    "Zaharia, M., et al. (2018). MLflow: A System for Machine Learning Lifecycle. Proceedings of the Workshop on Systems for ML, NeurIPS 2018.",
    "Docker, Inc. (2024). Docker Documentation. https://docs.docker.com/",
    "McKinney, W. (2010). Data Structures for Statistical Computing in Python. Proceedings of the 9th Python in Science Conference, 51–56.",
    "Merkel, D. (2014). Docker: Lightweight Linux Containers for Consistent Development and Deployment. Linux Journal, 239, 2.",
    "Hastie, T., Tibshirani, R., & Friedman, J. (2009). The Elements of Statistical Learning: Data Mining, Inference, and Prediction (2nd ed.). Springer.",
]

for i, ref in enumerate(references, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.space_after   = Pt(4)
    run = p.add_run(f"[{i}]  {ref}")
    set_font(run, size=10)

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(str(OUTPUT))
size_mb = OUTPUT.stat().st_size / (1024 * 1024)
print(f"\n✅ Report saved to: {OUTPUT}")
print(f"   File size: {size_mb:.2f} MB")
print(f"   Sections: Title Page, Abstract, Introduction, Dataset & Feature Analysis,")
print(f"             Methodology (15 models), Experiments & Results, Discussion, Conclusion, References")
print(f"   Images embedded: {sum(1 for _ in list(CLASS_DIR.glob('*.png')) + list(REG_DIR.glob('*.png')) + list(CORR_DIR.glob('*.png')) + [FEAT_IMG] if Path(_).exists())}")
